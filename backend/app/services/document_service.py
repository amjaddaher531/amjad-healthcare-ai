"""
Document ingestion utilities: turns uploaded files (PDF, DOCX, images,
scanned PDFs) into plain text, with OCR fallback for non-text content.
"""
from __future__ import annotations
import io
from dataclasses import dataclass

from pypdf import PdfReader
from PIL import Image
import pytesseract
from docx import Document as DocxDocument

from app.config import get_settings

_settings = get_settings()
pytesseract.pytesseract.tesseract_cmd = _settings.tesseract_cmd


@dataclass
class ExtractedFile:
    filename: str
    text: str
    page_count: int
    ocr_used: bool
    ocr_confidence: float | None


def _ocr_image_bytes(data: bytes) -> tuple[str, float]:
    image = Image.open(io.BytesIO(data))
    ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
    text = pytesseract.image_to_string(image)
    confidences = [int(c) for c in ocr_data.get("conf", []) if c not in ("-1", "", None) and int(c) >= 0]
    avg_conf = sum(confidences) / len(confidences) if confidences else 0.0
    return text, avg_conf


def extract_pdf(filename: str, data: bytes) -> ExtractedFile:
    reader = PdfReader(io.BytesIO(data))
    text_parts: list[str] = []
    ocr_used = False
    confidences: list[float] = []

    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            text_parts.append(page_text)
        else:
            # Scanned page with no extractable text layer -> rasterize + OCR
            try:
                from pdf2image import convert_from_bytes
                images = convert_from_bytes(data, first_page=reader.pages.index(page) + 1,
                                             last_page=reader.pages.index(page) + 1)
                if images:
                    buf = io.BytesIO()
                    images[0].save(buf, format="PNG")
                    ocr_text, conf = _ocr_image_bytes(buf.getvalue())
                    text_parts.append(ocr_text)
                    confidences.append(conf)
                    ocr_used = True
            except Exception:
                pass

    return ExtractedFile(
        filename=filename,
        text="\n".join(text_parts).strip(),
        page_count=len(reader.pages),
        ocr_used=ocr_used,
        ocr_confidence=(sum(confidences) / len(confidences)) if confidences else None,
    )


def extract_docx(filename: str, data: bytes) -> ExtractedFile:
    doc = DocxDocument(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " | ".join(cell.text for cell in row.cells)
    return ExtractedFile(filename=filename, text=text.strip(), page_count=1, ocr_used=False, ocr_confidence=None)


def extract_image(filename: str, data: bytes) -> ExtractedFile:
    text, conf = _ocr_image_bytes(data)
    return ExtractedFile(filename=filename, text=text.strip(), page_count=1, ocr_used=True, ocr_confidence=conf)


def extract_text_plain(filename: str, data: bytes) -> ExtractedFile:
    return ExtractedFile(filename=filename, text=data.decode("utf-8", errors="ignore"),
                          page_count=1, ocr_used=False, ocr_confidence=None)


def extract_file(filename: str, data: bytes) -> ExtractedFile:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_pdf(filename, data)
    if lower.endswith(".docx"):
        return extract_docx(filename, data)
    if lower.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp")):
        return extract_image(filename, data)
    if lower.endswith((".txt", ".csv")):
        return extract_text_plain(filename, data)
    raise ValueError(f"Unsupported file type: {filename}")
